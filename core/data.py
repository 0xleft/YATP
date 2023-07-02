import base64
import copy
import os

import json
from requests import post
from docx import Document
from docx.shared import Pt

from core.api import request_handler
from core.util import downscale_image
import eel


class ProblemImage:
    def __init__(self, image_path, description):
        self.image_path = image_path
        self.description = description

    def __str__(self):
        return f'Image path: {self.image_path}\nDescription: {self.description}'


# the converter class that holds references to all images and converts them later into a nice table in word
def make_images_smaller(path):
    try:
        os.mkdir(path + "/small")
    except FileExistsError:
        pass

    for file in os.listdir(path):
        if file.endswith(".jpg") or file.endswith(".png"):
            downscale_image(path + "/" + file, path + "/small/small_" + file, 700)


class Converter:
    def __init__(self):
        self.doc = Document()
        self.images = []

    def load_images(self, path):
        self.images = []
        for file in os.listdir(path):
            if file.endswith(".jpg") or file.endswith(".png"):
                try:
                    self.images.append(ProblemImage(path + "/" + file, ''))
                except Exception as e:
                    eel.show_notification(f"Error loading image {file}: {e}")

        print(self.images)

    def get_image_src(self, index):
        with open(self.images[index].image_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
        image_format = "png"
        if self.images[index].image_path.endswith(".jpg"):
            image_format = "jpeg"
        return f"data:image/{image_format};base64,{encoded_string}"

    def open_doc(self):
        self.doc.save('FILE_GENERATED_BY_YATP.docx')
        os.system('start FILE_GENERATED_BY_YATP.docx')

    def make_executive_summary(self, api_key):
        photos = ""
        for i in range(len(self.images)):
            if self.images[i].description == '':
                continue
            photos += f"P{i + 1}: {self.images[i].description} \n"

        if photos == "":
            eel.show_notification("None of the photos have descriptions")
            return
        print(photos)

        response = request_handler.model_complete_request_prompt(api_key, photos)

        if response.status_code != 200:
            eel.show_notification("Error",
                                  "Something went wrong while generating the executive summary. " + response.text)
            print(response.text)
            return

        response = response.json()
        print(response)
        response = response['choices'][0]['message']['content']
        self.doc.add_paragraph('')
        self.doc.add_paragraph('')
        self.doc.add_paragraph('Recommended executive summary: ')
        self.doc.add_paragraph(response)
        self.doc.add_paragraph('')
        self.doc.add_paragraph('')

    def convert_to_doc(self):
        def add_cell_data(cell, index, image_path, description):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_path, width=Pt(300))

            cell_paragraph = cell.add_paragraph()
            cell_paragraph.add_run(f"Foto {index + 1}: {description}").bold = True

            cell_paragraph.style.font.size = Pt(12)

        # we dont want to be editing the original list
        local_images = copy.deepcopy(self.images)

        try:
            os.remove('FILE_GENERATED_BY_YATP.docx')
        except FileNotFoundError:
            pass
        except PermissionError:
            eel.show_notification("Error", "Please close the file FILE_GENERATED_BY_YATP.docx")
            return

        self.doc = Document()

        to_remove = []
        for image in local_images:
            if image.description == '':
                print(f'No description for {image.image_path}')
                to_remove.append(image)

        # this is so it doesnt remove the images while iterating
        for image in to_remove:
            local_images.remove(image)

        # if no images
        if len(local_images) == 0:
            eel.show_notification("Error", "None of the images have descriptions")
            return

        print(len(local_images))
        for i in range(0, len(local_images), 2):

            table = self.doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'

            try:
                row_data = local_images[i], local_images[i + 1]
            except IndexError:
                row_data = local_images[i], None
            row = table.add_row().cells

            add_cell_data(row[0], i, row_data[0].image_path, row_data[0].description)
            if row_data[1] is not None:
                add_cell_data(row[1], i + 1, row_data[1].image_path, row_data[1].description)

    # save document to the folder that we are curently taking images from
    def save_doc(self, path):
        self.convert_to_doc()
        self.doc.save(path + '/FILE_GENERATED_BY_YATP.docx')

    # update description of the image
    def update_description(self, image_id, new_description):
        self.images[image_id].description = new_description
        print(self.images[image_id].description)
